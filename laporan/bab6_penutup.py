"""laporan/bab6_penutup.py - Bab 6 Penutup + Job Desk + Daftar Pustaka"""
from .helpers import heading, para, separator, add_table
from .config import ANGGOTA, PROYEK, MATKUL, KELAS, TANGGAL, EVALUASI


def tulis(doc):
    # ── BAB 6 PENUTUP ──────────────────────────────────────────────────────
    heading(doc, "BAB VI PENUTUP")
    separator(doc)

    heading(doc, "6.1 Kesimpulan", level=2)
    para(doc,
        f"Berdasarkan hasil perancangan, implementasi, dan pengujian yang telah dilakukan, "
        f"dapat disimpulkan bahwa proyek {PROYEK} berhasil dikembangkan sebagai aplikasi "
        f"manajemen screen time berbasis web menggunakan Java Spring Boot. Berikut adalah "
        f"poin-poin kesimpulan dari proyek ini:")

    kesimpulan = [
        "Konsep Pemrograman Berorientasi Objek berhasil diterapkan melalui lima kelas utama "
        "(User, AktivitasDigital, TopUp, Notifikasi, LaporanHarian) dengan relasi yang terstruktur "
        "meliputi komposisi, asosiasi, dan dependency.",
        "Antarmuka web yang terdiri dari delapan halaman berhasil diimplementasikan menggunakan "
        "Thymeleaf dan Spring MVC, memungkinkan pengguna berinteraksi dengan sistem secara intuitif.",
        "Basis data MySQL dengan lima tabel dan relasi foreign key berhasil digunakan untuk "
        "menyimpan dan mengelola data pengguna secara persisten.",
        "Logika perhitungan skor kesehatan digital berjalan sesuai formula yang dirancang, "
        "menghasilkan skor yang akurat berdasarkan total screen time dan jumlah pelanggaran batas.",
        "Fitur notifikasi dan laporan harian berhasil diimplementasikan, memberikan umpan balik "
        "informatif kepada pengguna tentang kebiasaan digital mereka.",
    ]
    for i, k in enumerate(kesimpulan, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(k).font.size = __import__('docx.shared', fromlist=['Pt']).Pt(11)

    heading(doc, "6.2 Saran", level=2)
    saran = [
        "Menambahkan enkripsi password menggunakan BCryptPasswordEncoder untuk meningkatkan keamanan akun.",
        "Mengimplementasikan fitur grafik visualisasi tren screen time menggunakan library Chart.js.",
        "Menambahkan fitur push notification real-time berbasis WebSocket atau Firebase Cloud Messaging.",
        "Melakukan deployment aplikasi ke cloud platform seperti Railway atau Heroku agar dapat diakses publik.",
        "Mengembangkan fitur analitik mingguan dan bulanan untuk memberikan insight jangka panjang kepada pengguna.",
    ]
    for s in saran:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(s).font.size = __import__('docx.shared', fromlist=['Pt']).Pt(11)

    doc.add_page_break()

    # ── JOB DESK ───────────────────────────────────────────────────────────
    heading(doc, "LAMPIRAN A – Job Desk Anggota Kelompok")
    separator(doc)
    para(doc,
        f"Berikut adalah pembagian tugas masing-masing anggota kelompok dalam pengerjaan "
        f"proyek {PROYEK} pada mata kuliah {MATKUL} Kelas {KELAS}, {EVALUASI}:")
    add_table(doc,
        ["No", "Nama Anggota", "Job Desk", "Detail Tugas"],
        [[str(a["no"]), a["nama"] or "—", a["jobdesk"] or "—", a["detail"] or "—"]
         for a in ANGGOTA],
        col_widths=[0.4, 1.8, 1.8, 2.8]
    )

    doc.add_page_break()

    # ── LAMPIRAN B – CARA MENJALANKAN ──────────────────────────────────────
    heading(doc, "LAMPIRAN B – Cara Menjalankan Aplikasi")
    separator(doc)

    heading(doc, "Persyaratan Sistem", level=2)
    add_table(doc,
        ["Komponen", "Versi Minimum", "Keterangan"],
        [
            ["Java JDK", "17",   "OpenJDK atau Oracle JDK"],
            ["MySQL",    "8.0",  "Harus berjalan sebelum aplikasi dijalankan"],
            ["Maven",    "3.8",  "Sudah termasuk dalam wrapper (mvnw)"],
            ["Browser",  "Modern", "Chrome, Firefox, Edge (versi terbaru)"],
        ],
        col_widths=[1.5, 1.5, 3.8]
    )

    heading(doc, "Langkah Menjalankan", level=2)
    steps = [
        ("Buat database",
         "mysql -u root -p < ETS/tubes_a11/demo/src/main/resources/database.sql"),
        ("Konfigurasi koneksi database",
         "Edit file: ETS/tubes_a11/demo/src/main/resources/application.properties\n"
         "spring.datasource.url=jdbc:mysql://localhost:3306/mindfull_db\n"
         "spring.datasource.username=root\n"
         "spring.datasource.password=YOUR_PASSWORD"),
        ("Jalankan aplikasi",
         "cd ETS/tubes_a11/demo\n./mvnw spring-boot:run"),
        ("Akses aplikasi di browser",
         "http://localhost:8080\nLogin dengan: username=admin, password=12345"),
    ]
    for i, (label, cmd) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {label}").bold = True
        from .helpers import code_block as cb
        cb(doc, cmd)

    doc.add_page_break()

    # ── DAFTAR PUSTAKA ─────────────────────────────────────────────────────
    heading(doc, "DAFTAR PUSTAKA")
    separator(doc)

    pustaka = [
        ("American Academy of Pediatrics (AAP).", "2023.",
         "Screen Time and Children.", "Diakses dari https://www.aap.org"),
        ("Deitel, P., & Deitel, H.", "2020.",
         "Java How to Program, Early Objects (11th ed.).", "Pearson Education."),
        ("Spring Framework Team.", "2024.",
         "Spring Boot Reference Documentation.",
         "Diakses dari https://docs.spring.io/spring-boot/docs/current/reference/html/"),
        ("Thymeleaf Team.", "2024.",
         "Thymeleaf 3.1 Tutorial: Using Thymeleaf.",
         "Diakses dari https://www.thymeleaf.org/doc/tutorials/3.1/usingthymeleaf.html"),
        ("Walls, C.", "2022.",
         "Spring in Action (6th ed.).", "Manning Publications."),
        ("We Are Social & Hootsuite.", "2024.",
         "Digital 2024: Indonesia.",
         "Diakses dari https://datareportal.com/reports/digital-2024-indonesia"),
        ("MySQL AB.", "2024.",
         "MySQL 8.0 Reference Manual.",
         "Diakses dari https://dev.mysql.com/doc/refman/8.0/en/"),
    ]

    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for (pengarang, tahun, judul, penerbit) in pustaka:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(28)
        p.paragraph_format.space_after = Pt(6)
        r1 = p.add_run(f"{pengarang} ({tahun}). ")
        r1.font.size = Pt(11)
        r2 = p.add_run(judul)
        r2.font.size = Pt(11)
        r2.italic = True
        r3 = p.add_run(f" {penerbit}")
        r3.font.size = Pt(11)
