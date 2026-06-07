"""laporan/helpers.py - Helper functions untuk formatting docx"""
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .config import C_NAVY, C_WHITE, C_BLACK, C_TEAL, C_BLUE, C_DKCODE


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if p.runs:
        run = p.runs[0]
        if level == 1:
            run.font.color.rgb = C_NAVY
            run.font.size = Pt(14)
        elif level == 2:
            run.font.color.rgb = C_TEAL
            run.font.size = Pt(12)
        else:
            run.font.color.rgb = C_BLUE
            run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def para(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    run.font.color.rgb = C_BLACK
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Pt(28)
    return p


def code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Pt(20)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(4)
    p.paragraph_format.first_line_indent = Pt(0)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1E2235')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name  = 'Consolas'
    run.font.size  = Pt(9)
    run.font.color.rgb = C_DKCODE
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, '0F3460')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = C_WHITE
        run.font.size = Pt(10)
    # Rows
    for ri, row_data in enumerate(rows):
        bg = 'EFF6FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = C_BLACK
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return table


def separator(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '0F3460')
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
